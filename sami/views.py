from django.shortcuts import render, redirect
from .models import Human
import os

from  django.contrib import messages
from django.contrib.auth.decorators import login_required

from django.http import HttpResponse
from docx import Document

from docx.shared import Inches
from io import BytesIO


def download_info(request, id):
    human_instance = Human.objects.get(id=id)

    # Create a DOCX document
    doc = Document()
    doc.add_heading('CV', level=1)

    # Add profile details to the document
    doc.add_paragraph(f"Name: {human_instance.name}")
    doc.add_paragraph(f"Email: {human_instance.email}")
    doc.add_paragraph(f"Age: {human_instance.age}")
    doc.add_paragraph(f"Gender: {human_instance.get_gender_display()}")
    doc.add_paragraph(f"Address: {human_instance.address}")
    doc.add_paragraph(f"Education: {human_instance.education}")
    doc.add_paragraph(f"Workplace: {human_instance.workplace}")
    doc.add_paragraph(f"Skills: {human_instance.skills}")
    doc.add_paragraph(f"Father's Name: {human_instance.father_name}")
    doc.add_paragraph(f"Mother's Name: {human_instance.mother_name}")

    # Add profile image to the document
    if human_instance.image:
        image_path = human_instance.image.path
        if image_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            doc.add_picture(image_path, width=Inches(2.0))

    # Set response headers for the file download
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={human_instance.name}_profile.docx'

    # Save the document to the response
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response.write(buffer.getvalue())
    return response



def hello(request):
    if not request.user.is_authenticated:
        messages.warning(request, 'Login first!')
        return redirect('login_user')
    
    if request.method == 'POST':
        name = request.POST.get('name')
        email = request.POST.get('email')
        age = request.POST.get('age')
        gender = request.POST.get('gender')
        image = request.FILES.get('image')
        address = request.POST.get('address')
        education = request.POST.get('education')
        workplace = request.POST.get('workplace')
        skills = request.POST.get('skills')
        father_name = request.POST.get('father_name')
        mother_name = request.POST.get('mother_name')

        human = Human.objects.create(
            name=name, email=email, age=age, gender=gender, image=image,
            address=address, education=education, workplace=workplace,
            skills=skills, father_name=father_name, mother_name=mother_name
        )

        human.save()

    return render(request, 'sami_template/index.html')

@login_required(login_url= 'login_user')
def all_prof(request):
    human = Human.objects.all()
    
    return render(request, 'sami_template/all_profile.html', locals())

def delete_prof(request, id):
    prof = Human.objects.get(id=id)
    if prof.image != 'def.jpg':
        os.remove(prof.image.path)
    prof.delete()
    return redirect('all_prof')

def update_prof(request, id):
    human_instance = Human.objects.get(id=id)
    if request.method == 'POST':
        name = request.POST.get('name')
        email = request.POST.get('email')
        age = request.POST.get('age')
        gender = request.POST.get('gender')
        image = request.FILES.get('image')
        address = request.POST.get('address')
        education = request.POST.get('education')
        workplace = request.POST.get('workplace')
        skills = request.POST.get('skills')
        father_name = request.POST.get('father_name')
        mother_name = request.POST.get('mother_name')

        if image:
            if human_instance.image != 'def.jpg':
                os.remove(human_instance.image.path)
            human_instance.image = image

        human_instance.name = name
        human_instance.email = email
        human_instance.age = age
        human_instance.gender = gender
        human_instance.address = address
        human_instance.education = education
        human_instance.workplace = workplace
        human_instance.skills = skills
        human_instance.father_name = father_name
        human_instance.mother_name = mother_name

        human_instance.save()
        return redirect('all_prof')

    return render(request, 'sami_template/update_prof.html', {'human_instance': human_instance})




from django.shortcuts import render, redirect
from .models import Task
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.views import View

@method_decorator(login_required, name='dispatch')
class TaskView(View):
    template_name = 'sami_template/tasks.html'

    def get(self, request):
        tasks = Task.objects.filter(user=request.user)
        return render(request, self.template_name, {'tasks': tasks})

    def post(self, request):
        title = request.POST.get('title')
        description = request.POST.get('description')
        date = request.POST.get('date')
        
        Task.objects.create(user=request.user, title=title, description=description, date=date)
        return redirect('tasks')

@login_required
def delete_task(request, id):
    task = Task.objects.get(id=id, user=request.user)
    task.delete()
    return redirect('tasks')

@login_required
def update_task(request, id):
    task = Task.objects.get(id=id, user=request.user)
    
    if request.method == 'POST':
        task.title = request.POST.get('title')
        task.description = request.POST.get('description')
        task.date = request.POST.get('date')
        task.save()
        return redirect('tasks')

    return render(request, 'sami_template/update_task.html', {'task': task})

